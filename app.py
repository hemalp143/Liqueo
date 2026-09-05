"""Streamlit web UI for Liqueo - Knowledge Discovery & Reuse System."""

import streamlit as st
from streamlit_option_menu import option_menu
import pandas as pd
from datetime import datetime
import json
import os
from pathlib import Path
import requests
from io import BytesIO

from liqueo.core import Document, KnowledgeBase, generate_doc_id
from liqueo.embeddings import EmbeddingsManager
from liqueo.recommender import RecommendationEngine
from liqueo.synthesizer import KnowledgeSynthesizer
from liqueo.workflow import WorkflowEngine, WorkflowSession


# File parsing functions
def extract_text_from_pdf(file):
    """Extract text from PDF file."""
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text[:5000]  # Limit to 5000 chars
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file):
    """Extract text from Word document."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text[:5000]  # Limit to 5000 chars
    except Exception as e:
        return f"Error reading Word document: {str(e)}"

def extract_text_from_xlsx(file):
    """Extract text from Excel file."""
    try:
        df = pd.read_excel(file)
        text = df.to_string()
        return text[:5000]  # Limit to 5000 chars
    except Exception as e:
        return f"Error reading Excel file: {str(e)}"

def extract_text_from_txt(file):
    """Extract text from text file."""
    try:
        text = file.read().decode('utf-8')
        return text[:5000]
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def extract_text_from_csv(file):
    """Extract text from CSV file."""
    try:
        df = pd.read_csv(file)
        text = df.to_string()
        return text[:5000]
    except Exception as e:
        return f"Error reading CSV file: {str(e)}"

def parse_uploaded_file(file):
    """Parse uploaded file and extract content."""
    filename = file.name.lower()

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        return extract_text_from_docx(file)
    elif filename.endswith('.xlsx') or filename.endswith('.xls'):
        return extract_text_from_xlsx(file)
    elif filename.endswith('.txt'):
        return extract_text_from_txt(file)
    elif filename.endswith('.csv'):
        return extract_text_from_csv(file)
    else:
        return f"Unsupported file type: {filename}"

def download_file_from_url(url):
    """Download and parse file from URL (OneDrive, SharePoint, etc.)."""
    try:
        # For OneDrive links, modify URL to get direct download
        if 'onedrive.live.com' in url or 'sharepoint.com' in url:
            # Remove ?download=1 and add it
            url = url.split('?')[0]
            if '?download=1' not in url:
                url = url + '?download=1'

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        # Get filename from URL or use generic name
        filename = url.split('/')[-1].split('?')[0] or "downloaded_file"

        # Create a file-like object
        file_content = BytesIO(response.content)
        file_content.name = filename

        return file_content, filename
    except Exception as e:
        return None, f"Error downloading file: {str(e)}"


# Configure page
st.set_page_config(
    page_title="Liqueo - Knowledge Discovery",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .metric-card {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        margin: 10px 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 15px;
        border-radius: 5px;
        border-left: 4px solid #28a745;
    }
    .info-box {
        background-color: #d1ecf1;
        padding: 15px;
        border-radius: 5px;
        border-left: 4px solid #17a2b8;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 15px;
        border-radius: 5px;
        border-left: 4px solid #ffc107;
    }
</style>
""", unsafe_allow_html=True)


# Initialize session state
if "kb" not in st.session_state:
    st.session_state.kb = KnowledgeBase()
    st.session_state.embeddings = EmbeddingsManager(st.session_state.kb)
    st.session_state.recommender = RecommendationEngine(st.session_state.embeddings)
    st.session_state.synthesizer = KnowledgeSynthesizer(
        st.session_state.kb, st.session_state.recommender
    )
    st.session_state.workflow_engine = WorkflowEngine()

# Initialize modal state
if "show_modal" not in st.session_state:
    st.session_state.show_modal = False
if "selected_doc" not in st.session_state:
    st.session_state.selected_doc = None

# Initialize workflow state
if "current_workflow" not in st.session_state:
    st.session_state.current_workflow = None


def render_header():
    """Render page header."""
    col1, col2 = st.columns([3, 1])
    with col1:
        st.title("🧠 Liqueo")
        st.markdown("**Knowledge Discovery & Reuse for Financial Consultants**")
    with col2:
        st.info("v1.0.0")


def render_home():
    """Render home/dashboard page."""
    render_header()

    st.markdown("---")

    # Dashboard metrics
    col1, col2, col3, col4 = st.columns(4)

    docs = st.session_state.kb.list_documents()
    industries = set(d.industry for d in docs if d.industry)
    types = set(d.transaction_type for d in docs if d.transaction_type)
    total_value = sum(d.engagement_value or 0 for d in docs)

    with col1:
        st.metric("Total Documents", len(docs))
    with col2:
        st.metric("Industries", len(industries))
    with col3:
        st.metric("Transaction Types", len(types))
    with col4:
        st.metric("Total Value ($M)", f"${total_value:.1f}")

    st.markdown("---")

    # Key features
    st.subheader("Key Features")

    feature_cols = st.columns(3)

    with feature_cols[0]:
        st.info("""
        ### 📚 Knowledge Base
        Store and organize consulting engagements with rich metadata
        """)

    with feature_cols[1]:
        st.success("""
        ### 🔍 Semantic Search
        Find similar cases using embeddings or keyword matching
        """)

    with feature_cols[2]:
        st.warning("""
        ### 💡 AI Synthesis
        Generate insights from similar engagements
        """)

    st.markdown("---")

    # Sample data
    if len(docs) == 0:
        st.info("👉 Get started by adding documents from the sidebar!")
    else:
        st.subheader("Recent Engagements")

        df_data = []
        for doc in sorted(docs, key=lambda d: d.created_at, reverse=True)[:5]:
            df_data.append({
                "Title": doc.title[:50],
                "Industry": doc.industry,
                "Type": doc.transaction_type,
                "Value ($M)": doc.engagement_value,
                "Duration (months)": doc.duration_months
            })

        df = pd.DataFrame(df_data)
        st.dataframe(df, use_container_width=True, hide_index=True)


def render_add_document():
    """Render add document page."""
    render_header()
    st.subheader("Add New Engagement")

    st.markdown("---")

    with st.form("add_doc_form"):
        st.markdown("### Engagement Details")
        col1, col2 = st.columns(2)

        with col1:
            title = st.text_input("Engagement Title *", placeholder="e.g., SaaS Acquisition")
            industry = st.text_input("Industry *", placeholder="e.g., Technology")
            transaction_type = st.text_input("Transaction Type", placeholder="e.g., M&A, Restructuring")

        with col2:
            value = st.number_input("Engagement Value ($M)", min_value=0.0, step=0.1)
            duration = st.number_input("Duration (months)", min_value=1, max_value=60)
            client = st.text_input("Client Name", placeholder="e.g., ABC Corporation")

        st.markdown("---")
        st.markdown("### Add Content")
        st.markdown("*Choose one: Type details below OR upload files/URLs*")

        # File upload options
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("**📄 Upload Local Files**")
            st.markdown("Supported: PDF, Word, Excel, CSV, Text")
            uploaded_files = st.file_uploader(
                "Select files",
                type=["pdf", "docx", "doc", "xlsx", "xls", "csv", "txt"],
                accept_multiple_files=True,
                key="manual_file_uploader"
            )

        with col2:
            st.markdown("**🔗 Download from URL**")
            st.markdown("OneDrive, SharePoint, etc.")
            file_url = st.text_input(
                "Paste URL",
                placeholder="https://onedrive.live.com/...",
                key="manual_url_input",
                help="Direct download link to document"
            )

        # Process uploaded and downloaded files
        extracted_content = ""
        file_list = []
        local_files = []
        url_files = []

        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} local file(s) selected")
            for file in uploaded_files:
                extracted = parse_uploaded_file(file)
                extracted_content += f"\n--- From {file.name} ---\n{extracted}"
                file_list.append(file.name)
                local_files.append(file.name)

        if file_url:
            if st.form_submit_button("📥 Download from URL", use_container_width=False):
                with st.spinner("Downloading..."):
                    file_content, filename = download_file_from_url(file_url)
                    if file_content is None:
                        st.error(f"❌ {filename}")
                    else:
                        st.success(f"✅ Downloaded: {filename}")
                        extracted = parse_uploaded_file(file_content)
                        extracted_content += f"\n--- From {filename} ---\n{extracted}"
                        file_list.append(filename)
                        url_files.append(filename)
                        st.rerun()

        st.markdown("---")

        # Content textarea (can be auto-filled or edited)
        default_content = extracted_content if extracted_content else ""
        content = st.text_area(
            "Engagement Details *",
            height=150,
            placeholder="Type engagement details OR upload files above...",
            value=default_content
        )

        outcomes = st.text_area(
            "Key Outcomes",
            height=80,
            placeholder="Summary of outcomes and impact..."
        )

        approach = st.text_area(
            "Consulting Approach",
            height=80,
            placeholder="Methodology and key approach..."
        )

        if st.form_submit_button("✅ Add Engagement", use_container_width=True):
            if not title or not industry or not content:
                st.error("Please fill in Title, Industry, and Details (marked with *)")
            else:
                doc_id = generate_doc_id(title, datetime.now())
                metadata = {}
                if file_list:
                    metadata = {
                        "uploaded_files": file_list,
                        "file_count": len(file_list),
                        "local_files": local_files,
                        "url_files": url_files
                    }

                doc = Document(
                    id=doc_id,
                    title=title,
                    content=content,
                    doc_type="engagement",
                    industry=industry,
                    transaction_type=transaction_type or None,
                    engagement_value=value if value > 0 else None,
                    duration_months=duration,
                    client_name=client or None,
                    key_outcomes=outcomes or None,
                    consulting_approach=approach or None,
                    metadata=metadata if metadata else None
                )

                st.session_state.kb.add_document(doc)

                # Try to generate embeddings
                try:
                    st.session_state.embeddings.embed_document(doc)
                    msg = f"✓ Document added: **{title}**\n\nEmbeddings generated successfully"
                    if file_list:
                        msg += f"\n\n📁 Files: {', '.join(file_list)}"
                    st.success(msg)
                except Exception as e:
                    st.warning(f"✓ Document added: **{title}**\n\n⚠️ Embeddings not available (using keyword search instead)")



def render_search():
    """Render search page."""
    render_header()
    st.subheader("Search Knowledge Base")

    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        query = st.text_input("Search Query", placeholder="e.g., cloud infrastructure optimization")
    with col2:
        top_k = st.slider("Top Results", 1, 10, 5)
    with col3:
        industry_filter = st.text_input("Filter by Industry", placeholder="Leave blank for no filter")

    if st.button("Search", use_container_width=True):
        if not query:
            st.warning("Please enter a search query")
        else:
            with st.spinner("Searching..."):
                filters = {}
                if industry_filter:
                    filters["industry"] = industry_filter

                results = st.session_state.embeddings.semantic_search(
                    query, top_k=top_k, filters=filters
                )

                if not results:
                    st.info("No results found")
                else:
                    st.success(f"Found {len(results)} matching engagement(s)")

                    for i, result in enumerate(results, 1):
                        doc = result.document
                        with st.container(border=True):
                            col1, col2 = st.columns([3, 1])

                            with col1:
                                st.markdown(f"**{i}. {doc.title}**")
                                st.markdown(f"*{doc.industry} • {doc.transaction_type}*")

                            with col2:
                                st.metric("Relevance", f"{result.similarity_score:.0%}")

                            col1, col2, col3, col4 = st.columns(4)
                            with col1:
                                st.caption(f"Value: ${doc.engagement_value}M" if doc.engagement_value else "Value: N/A")
                            with col2:
                                st.caption(f"Duration: {doc.duration_months}m" if doc.duration_months else "Duration: N/A")
                            with col3:
                                st.caption(f"Client: {doc.client_name}" if doc.client_name else "Client: N/A")
                            with col4:
                                if st.button(f"👁️ View", key=f"view_{doc.id}", use_container_width=True):
                                    st.session_state.selected_doc = doc
                                    st.session_state.show_modal = True

                            if doc.key_outcomes:
                                st.markdown(f"**Outcomes:** {doc.key_outcomes}")

    # Display modal popup if a document is selected
    if st.session_state.get("show_modal") and st.session_state.get("selected_doc"):
        doc = st.session_state.selected_doc
        with st.modal("📋 Engagement Details"):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.title(doc.title)
            with col2:
                if st.button("✕", key="close_modal", help="Close"):
                    st.session_state.show_modal = False
                    st.rerun()

            st.markdown("---")

            # Key Information
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Industry", doc.industry or "N/A")
            with col2:
                st.metric("Type", doc.transaction_type or "N/A")
            with col3:
                st.metric("Value ($M)", f"${doc.engagement_value}" if doc.engagement_value else "N/A")
            with col4:
                st.metric("Duration (months)", doc.duration_months or "N/A")

            st.markdown("---")

            # Client Information
            if doc.client_name:
                st.subheader("Client")
                st.write(doc.client_name)

            # Engagement Details
            st.subheader("📝 Engagement Details")
            st.write(doc.content)

            # Consulting Approach
            if doc.consulting_approach:
                st.subheader("🎯 Consulting Approach")
                st.write(doc.consulting_approach)

            # Key Outcomes
            if doc.key_outcomes:
                st.subheader("✅ Key Outcomes")
                st.write(doc.key_outcomes)

            # Metadata
            if doc.metadata:
                st.subheader("📁 File Information")
                if doc.metadata.get("uploaded_files"):
                    st.write("**Uploaded Files:**")
                    for fname in doc.metadata.get("uploaded_files", []):
                        st.caption(f"• {fname}")

            st.markdown("---")
            col1, col2 = st.columns(2)
            with col1:
                st.caption(f"📅 Created: {doc.created_at.strftime('%Y-%m-%d %H:%M') if doc.created_at else 'N/A'}")
            with col2:
                st.caption(f"🆔 ID: {doc.id[:8]}...")


def render_recommendations():
    """Render recommendations page."""
    render_header()
    st.subheader("Get Recommendations")

    st.markdown("---")

    engagement = st.text_area(
        "Describe the Current Engagement",
        height=150,
        placeholder="Describe the client situation, challenges, and desired outcomes...",
        value="We need to optimize our cloud infrastructure and reduce operational costs"
    )

    top_k = st.slider("Number of Recommendations", 1, 10, 3)

    if st.button("Get Recommendations", use_container_width=True):
        if not engagement:
            st.warning("Please describe the engagement")
        else:
            with st.spinner("Analyzing similar engagements..."):
                recommendations = st.session_state.recommender.recommend_similar_engagements(
                    engagement, top_k=top_k
                )

                if not recommendations:
                    st.info("No recommendations found")
                else:
                    st.success(f"Found {len(recommendations)} similar engagement(s)")

                    for i, rec in enumerate(recommendations, 1):
                        doc = rec.reference_document
                        with st.container(border=True):
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.markdown(f"**{i}. {doc.title}**")
                                st.markdown(f"*Relevance: {rec.relevance_score:.0%} • {rec.reasoning}*")
                            with col2:
                                if st.button(f"👁️ View", key=f"view_rec_{doc.id}", use_container_width=True):
                                    st.session_state.selected_doc = doc
                                    st.session_state.show_modal = True

                            col1, col2, col3 = st.columns(3)
                            with col1:
                                if rec.suggested_approach:
                                    st.info(f"**Approach:** {rec.suggested_approach}")
                            with col2:
                                if rec.potential_challenges:
                                    st.warning(f"**Challenges:** {rec.potential_challenges[:100]}...")
                            with col3:
                                if rec.estimated_effort:
                                    st.success(f"**Effort:** {rec.estimated_effort}")

    # Display modal popup if a document is selected
    if st.session_state.get("show_modal") and st.session_state.get("selected_doc"):
        doc = st.session_state.selected_doc
        with st.modal("📋 Engagement Details"):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.title(doc.title)
            with col2:
                if st.button("✕", key="close_modal_rec", help="Close"):
                    st.session_state.show_modal = False
                    st.rerun()

            st.markdown("---")

            # Key Information
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Industry", doc.industry or "N/A")
            with col2:
                st.metric("Type", doc.transaction_type or "N/A")
            with col3:
                st.metric("Value ($M)", f"${doc.engagement_value}" if doc.engagement_value else "N/A")
            with col4:
                st.metric("Duration (months)", doc.duration_months or "N/A")

            st.markdown("---")

            # Client Information
            if doc.client_name:
                st.subheader("Client")
                st.write(doc.client_name)

            # Engagement Details
            st.subheader("📝 Engagement Details")
            st.write(doc.content)

            # Consulting Approach
            if doc.consulting_approach:
                st.subheader("🎯 Consulting Approach")
                st.write(doc.consulting_approach)

            # Key Outcomes
            if doc.key_outcomes:
                st.subheader("✅ Key Outcomes")
                st.write(doc.key_outcomes)

            # Metadata
            if doc.metadata:
                st.subheader("📁 File Information")
                if doc.metadata.get("uploaded_files"):
                    st.write("**Uploaded Files:**")
                    for fname in doc.metadata.get("uploaded_files", []):
                        st.caption(f"• {fname}")

            st.markdown("---")
            col1, col2 = st.columns(2)
            with col1:
                st.caption(f"📅 Created: {doc.created_at.strftime('%Y-%m-%d %H:%M') if doc.created_at else 'N/A'}")
            with col2:
                st.caption(f"🆔 ID: {doc.id[:8]}...")


def render_synthesize():
    """Render synthesis page."""
    render_header()
    st.subheader("Synthesize Insights")

    st.markdown("---")

    engagement = st.text_area(
        "Engagement Description",
        height=150,
        placeholder="Describe the engagement and what insights you need...",
        value="We need to optimize our technology infrastructure and reduce operational costs"
    )

    top_k = st.slider("Consider Top Similar Cases", 1, 10, 3)

    if st.button("Generate Insights", use_container_width=True):
        if not engagement:
            st.warning("Please describe the engagement")
        else:
            with st.spinner("Generating synthesis and insights..."):
                synthesis = st.session_state.synthesizer.synthesize_recommendations(
                    engagement, top_k=top_k
                )

                st.markdown("---")
                st.markdown(synthesis)


def render_analysis():
    """Render industry analysis page."""
    render_header()
    st.subheader("Industry Analysis")

    st.markdown("---")

    docs = st.session_state.kb.list_documents()
    industries = sorted(set(d.industry for d in docs if d.industry))

    if not industries:
        st.info("Add documents first to analyze industries")
    else:
        industry = st.selectbox("Select Industry", industries)

        if st.button("Analyze Trends", use_container_width=True):
            with st.spinner(f"Analyzing {industry} industry trends..."):
                analysis = st.session_state.synthesizer.analyze_industry_trends(industry)
                st.markdown(analysis)


def render_documents():
    """Render documents management page."""
    render_header()
    st.subheader("Knowledge Base Documents")

    st.markdown("---")

    docs = st.session_state.kb.list_documents()

    if not docs:
        st.info("No documents in knowledge base. Add documents to get started!")
    else:
        st.info(f"Total: {len(docs)} document(s)")

        # Create dataframe
        data = []
        for doc in sorted(docs, key=lambda d: d.created_at, reverse=True):
            data.append({
                "ID": doc.id[:12],
                "Title": doc.title,
                "Industry": doc.industry or "-",
                "Type": doc.transaction_type or "-",
                "Value": f"${doc.engagement_value}M" if doc.engagement_value else "-",
                "Duration": f"{doc.duration_months}m" if doc.duration_months else "-",
                "Created": doc.created_at.strftime("%Y-%m-%d")
            })

        df = pd.DataFrame(data)
        st.dataframe(df, use_container_width=True, hide_index=True)

        # Export option
        st.markdown("---")
        if st.button("Export Knowledge Base (JSON)", use_container_width=True):
            export_data = {
                "timestamp": datetime.now().isoformat(),
                "document_count": len(docs),
                "documents": [doc.to_dict() for doc in docs]
            }

            st.json(export_data)

            # Download button
            st.download_button(
                "Download JSON",
                json.dumps(export_data, indent=2, default=str),
                "liqueo_export.json"
            )


def render_workflow():
    """Render the 9-step knowledge discovery and reuse workflow."""
    render_header()
    st.subheader("Knowledge Discovery & Reuse Workflow")

    st.markdown("""
    ### 📚 Complete 9-Step Workflow
    Follow this guided workflow to discover relevant knowledge from past engagements
    and reuse it to create better outcomes for your current challenges.
    """)

    st.markdown("---")

    # Workflow creation
    col1, col2 = st.columns([2, 1])
    with col1:
        problem = st.text_area(
            "Step 1️⃣: Define Your Problem or Task *",
            height=100,
            placeholder="Describe the challenge, problem, or task you need to solve..."
        )
    with col2:
        st.info("**Step 1: Identify**\n\nDefine the task or problem you need to solve")

    if problem and st.button("Start Knowledge Workflow", use_container_width=True):
        workflow = st.session_state.workflow_engine.create_workflow(problem)
        st.session_state.current_workflow = workflow
        st.success("✅ Workflow created! Continue with Step 2...")
        st.rerun()

    st.markdown("---")

    # If workflow exists, show the steps
    if st.session_state.current_workflow:
        workflow = st.session_state.current_workflow
        progress = workflow.get_progress()

        # Progress bar
        st.markdown(f"### Progress: {progress['completed']}/{progress['total']} steps completed")
        st.progress(progress['percentage'] / 100)

        st.markdown("---")

        # Step 2: Search Knowledge
        with st.expander("Step 2️⃣: Search for Relevant Knowledge", expanded=progress['completed'] < 2):
            st.markdown("**Search for similar past engagements**")
            search_query = st.text_input(
                "Search query",
                placeholder="e.g., technology infrastructure optimization",
                key="workflow_search"
            )
            if st.button("Search Knowledge Base", key="workflow_search_btn"):
                if search_query:
                    results = st.session_state.embeddings.semantic_search(search_query, top_k=5)
                    st.session_state.workflow_search_results = results
                    st.success(f"Found {len(results)} similar engagements")

        # Step 3: Identify Related Documents
        with st.expander("Step 3️⃣: System Identifies Related Documents", expanded=progress['completed'] < 3):
            st.markdown("**AI-identified related documents, templates, and lessons learned**")
            if hasattr(st.session_state, 'workflow_search_results'):
                st.markdown("**Related Engagements:**")
                for i, result in enumerate(st.session_state.workflow_search_results, 1):
                    doc = result.document
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"{i}. **{doc.title}** ({result.similarity_score:.0%})")
                    with col2:
                        if st.checkbox("Select", key=f"select_{doc.id}"):
                            if doc.id not in workflow.selected_documents:
                                workflow.selected_documents.append(doc.id)

        # Step 4: AI Summarize & Recommend
        with st.expander("Step 4️⃣: AI Assists with Summaries & Recommendations", expanded=progress['completed'] < 4):
            st.markdown("**AI-powered summaries and strategic recommendations**")
            if workflow.selected_documents:
                st.success("✅ Analyzing selected documents...")
                for doc_id in workflow.selected_documents:
                    doc = st.session_state.kb.get_document(doc_id)
                    if doc:
                        with st.container(border=True):
                            st.markdown(f"**{doc.title}**")
                            st.markdown(f"*Industry: {doc.industry} | Type: {doc.transaction_type}*")
                            if doc.key_outcomes:
                                st.markdown(f"**Outcomes:** {doc.key_outcomes}")
                            if doc.consulting_approach:
                                st.markdown(f"**Approach:** {doc.consulting_approach}")
            else:
                st.info("Select documents in Step 3 to see summaries")

        # Step 5: Review & Evaluate
        with st.expander("Step 5️⃣: Review & Evaluate Results", expanded=progress['completed'] < 5):
            st.markdown("**Review the recommended content and assess relevance**")
            review_notes = st.text_area(
                "Your evaluation and notes",
                height=100,
                placeholder="What insights are most relevant? Any gaps?",
                key="workflow_review"
            )
            if st.button("Mark Step 5 Complete", key="step5_complete"):
                st.session_state.workflow_engine.update_step(
                    workflow.id, 5, "completed",
                    {"review_notes": review_notes}
                )
                st.success("✅ Step 5 completed")

        # Step 6: Select Content for Reuse
        with st.expander("Step 6️⃣: Select Content to Reuse/Adapt", expanded=progress['completed'] < 6):
            st.markdown("**Choose which elements to adapt for your solution**")
            elements = st.multiselect(
                "Select elements to reuse",
                ["Consulting Approach", "Success Factors", "Team Structure", "Timeline", "Process Steps", "Risk Mitigation"],
                key="workflow_reuse_elements"
            )
            if st.button("Mark Step 6 Complete", key="step6_complete"):
                st.session_state.workflow_engine.update_step(
                    workflow.id, 6, "completed",
                    {"selected_elements": elements}
                )
                st.success("✅ Step 6 completed")

        # Step 7: Create New Output
        with st.expander("Step 7️⃣: Create New Output", expanded=progress['completed'] < 7):
            st.markdown("**Create a new engagement or proposal using discovered knowledge**")
            col1, col2 = st.columns(2)
            with col1:
                new_title = st.text_input("Output Title", placeholder="e.g., New Client Proposal")
                new_industry = st.text_input("Industry")
            with col2:
                new_type = st.text_input("Type")
                new_value = st.number_input("Value ($M)", min_value=0.0, step=0.1)

            new_content = st.text_area(
                "Output Content (auto-populated from selected engagements)",
                height=150,
                key="workflow_new_content"
            )

            if st.button("Save Output", key="step7_save"):
                if new_title and new_industry and new_content:
                    doc_id = generate_doc_id(new_title, datetime.now())
                    new_doc = Document(
                        id=doc_id,
                        title=new_title,
                        content=new_content,
                        doc_type="engagement",
                        industry=new_industry,
                        transaction_type=new_type or None,
                        engagement_value=new_value if new_value > 0 else None,
                        metadata={"created_from_workflow": True, "source_docs": workflow.selected_documents}
                    )
                    workflow.created_output = new_doc
                    st.session_state.workflow_engine.update_step(workflow.id, 7, "completed")
                    st.success("✅ Output created! Continue to Step 8...")

        # Step 8: Tag & Classify
        with st.expander("Step 8️⃣: Tag & Classify the Output", expanded=progress['completed'] < 8):
            st.markdown("**Add tags and classification for future discovery**")
            tags = st.multiselect(
                "Add tags",
                ["Cost Optimization", "Digital Transformation", "M&A", "Restructuring", "Cloud Migration", "Custom"],
                key="workflow_tags"
            )
            custom_tag = st.text_input("Custom tag (if needed)")
            if custom_tag:
                tags.append(custom_tag)

            if st.button("Mark Step 8 Complete", key="step8_complete"):
                workflow.tags = tags
                st.session_state.workflow_engine.update_step(workflow.id, 8, "completed", {"tags": tags})
                st.success("✅ Step 8 completed")

        # Step 9: Store for Future Reuse
        with st.expander("Step 9️⃣: Store for Future Discovery & Reuse", expanded=progress['completed'] < 9):
            st.markdown("**Your output is now part of the knowledge base for future reuse**")

            if workflow.created_output:
                st.info(f"""
                **Document stored:**
                - Title: {workflow.created_output.title}
                - Industry: {workflow.created_output.industry}
                - Tags: {', '.join(workflow.tags)}
                - ID: {workflow.created_output.id[:12]}...

                This knowledge is now discoverable by other team members!
                """)

                if st.button("Complete Workflow", key="workflow_complete"):
                    st.session_state.workflow_engine.update_step(workflow.id, 9, "completed")
                    insights = st.session_state.workflow_engine.get_workflow_insights(workflow.id)
                    st.success("✅ Workflow completed!")
                    st.json(insights)
            else:
                st.warning("Create output in Step 7 first")

        st.markdown("---")
        st.markdown("### Workflow Summary")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Steps Completed", progress['completed'])
        with col2:
            st.metric("Documents Selected", len(workflow.selected_documents))
        with col3:
            st.metric("Tags Applied", len(workflow.tags))
        with col4:
            st.metric("Output Created", "✅" if workflow.created_output else "⏳")


def main():
    """Main app entry point."""
    # Sidebar navigation
    with st.sidebar:
        st.markdown("## Navigation")
        selected = option_menu(
            menu_title=None,
            options=["Home", "Knowledge Workflow", "Search", "Recommendations", "Synthesize", "Analysis", "Documents", "Add Engagement"],
            icons=["house", "flow", "search", "lightbulb", "sparkles", "chart-line", "folder", "plus"],
            menu_icon="cast",
            default_index=0,
            styles={
                "container": {"padding": "0!important", "background-color": "#f0f2f6"},
                "icon": {"color": "orange", "font-size": "25px"},
                "nav-link": {"font-size": "16px", "text-align": "left", "margin": "0px", "--hover-color": "#eee"},
                "nav-link-selected": {"background-color": "#FF6B35"},
            }
        )

        st.markdown("---")

        # Info section
        st.markdown("### About Liqueo")
        st.markdown("""
        **Knowledge Discovery & Reuse System** for financial consultants.

        - 🧠 Semantic search with embeddings
        - 💡 AI-powered recommendations
        - 📊 Industry trend analysis
        - 📈 Strategic insights synthesis
        """)

        # Status
        docs_count = len(st.session_state.kb.list_documents())
        st.markdown(f"**Status:** {docs_count} documents loaded")

    # Route to selected page
    if selected == "Home":
        render_home()
    elif selected == "Knowledge Workflow":
        render_workflow()
    elif selected == "Add Engagement":
        render_add_document()
    elif selected == "Search":
        render_search()
    elif selected == "Recommendations":
        render_recommendations()
    elif selected == "Synthesize":
        render_synthesize()
    elif selected == "Analysis":
        render_analysis()
    elif selected == "Documents":
        render_documents()


if __name__ == "__main__":
    main()
